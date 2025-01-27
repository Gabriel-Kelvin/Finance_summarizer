import os
import pandas as pd
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI
from docx import Document
from dotenv import load_dotenv

load_dotenv()

DOCUMENT_INTELLIGENCE_ENDPOINT = os.getenv("DOCUMENTINTELLIGENCE_ENDPOINT")
DOCUMENT_INTELLIGENCE_API_KEY = os.getenv("DOCUMENTINTELLIGENCE_API_KEY")
document_intelligence_client = DocumentIntelligenceClient(
    endpoint=DOCUMENT_INTELLIGENCE_ENDPOINT,
    credential=AzureKeyCredential(DOCUMENT_INTELLIGENCE_API_KEY)
)

AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
DEPLOYMENT_NAME = "gpt-4o"
openai_client = AzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_version="2024-08-01-preview"
)

document_analysis_client = DocumentAnalysisClient(

        endpoint="https://neohack.cognitiveservices.azure.com/",

        credential=AzureKeyCredential(
            "D4Vd1Xljol6cS0q8PoZiQFXz8uH1yYeXCIyhTHU5vdBxBhbeX9JLJQQJ99ALACYeBjFXJ3w3AAALACOGIiVt")
    )

def process_financial_report(pdf_path):
    """
    Extracts financial data from a PDF using Azure Document Intelligence.
    """
    try:
        # with open(pdf_path, "rb") as f:
        #     poller = document_intelligence_client.begin_analyze_document(
        #         model_id="prebuilt-layout",
        #         document=f
        #     )
        #     result = poller.result()
        with open(pdf_path, "rb") as file:
            poller = document_analysis_client.begin_analyze_document("prebuilt-layout", document=file)

            result = poller.result()

        extracted_data = {}
        for page in result.pages:
            for table in page.tables:
                for cell in table.cells:
                    extracted_data[cell.content] = cell.bounding_regions[0].polygon if cell.bounding_regions else None
        return extracted_data

    except Exception as e:
        raise RuntimeError(f"Error processing financial report: {e}")

def calculate_ratios(data):
    """
    Calculates financial ratios from extracted financial data.
    """
    try:
        total_assets = float(data.get("Total Assets", 1))
        total_liabilities = float(data.get("Total Liabilities", 1))
        shareholder_equity = float(data.get("Shareholder Equity", 1))
        net_income = float(data.get("Net Income", 1))
        total_revenue = float(data.get("Total Revenue", 1))
        earnings_per_share = float(data.get("EPS", 1))

        ratios = {
            "Current Ratio": total_assets / total_liabilities,
            "Debt-to-Equity Ratio": total_liabilities / shareholder_equity,
            "Return on Equity (ROE)": (net_income / shareholder_equity) * 100,
            "Return on Assets (ROA)": (net_income / total_assets) * 100,
            "Gross Profit Margin": (net_income / total_revenue) * 100,
            "Net Profit Margin": (net_income / total_revenue) * 100,
            "Earnings Per Share (EPS)": earnings_per_share
        }
        return ratios

    except Exception as e:
        raise ValueError(f"Error calculating financial ratios: {e}")

def generate_summary(data, ratios):
    """
    Generates a financial performance summary using Azure OpenAI.
    """
    prompt = f"""
    Summarize the following financial performance data:
    Financial Data: {data}
    Ratios: {ratios}
    Include risk factors and overall performance highlights.
    """
    try:
        response = openai_client.Completion.create(
            engine=DEPLOYMENT_NAME,
            prompt=prompt,
            max_tokens=300,
            temperature=0.7
        )
        return response["choices"][0]["text"]

    except Exception as e:
        raise RuntimeError(f"Error generating summary: {e}")

def create_excel(ratios, output_path="financial_ratios.xlsx"):
    """
    Saves the calculated ratios into an Excel file.
    """
    try:
        df = pd.DataFrame([ratios])
        df.to_excel(output_path, index=False)
        print(f"Excel file saved at {output_path}")
    except Exception as e:
        raise RuntimeError(f"Error creating Excel file: {e}")

def create_word_document(summary, output_path="financial_summary.docx"):
    """
    Generates a Word document with the financial performance summary.
    """
    try:
        doc = Document()
        doc.add_heading("Financial Summary", level=1)
        doc.add_paragraph(summary)
        doc.save(output_path)
        print(f"Word document saved at {output_path}")
    except Exception as e:
        raise RuntimeError(f"Error creating Word document: {e}")

if __name__ == "__main__":
    pdf_path = "Report.pdf"

    try:
        extracted_data = process_financial_report(pdf_path)

        ratios = calculate_ratios(extracted_data)

        summary = generate_summary(extracted_data, ratios)

        create_excel(ratios)

        create_word_document(summary)

        print("Financial analysis completed successfully!")

    except Exception as e:
        print(f"Error: {e}")

