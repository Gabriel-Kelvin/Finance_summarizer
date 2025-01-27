import os
from openai import AzureOpenAI
import pandas as pd
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
from dotenv import load_dotenv
from docx import Document
import collections
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest

collections.Callable = collections.abc.Callable

load_dotenv()

client = AzureOpenAI(

    api_key="3txjyxDMw3BbClF4YSqpHwGPsOfoSlRFl22zkrvU2KiIv6tcFdwcJQQJ99ALAC77bzfXJ3w3AAABACOG41Xj",

    api_version="2024-08-01-preview",

    azure_endpoint="https://neohackathon01-2024.openai.azure.com/"

)

document_intelligence_client = DocumentIntelligenceClient(
    endpoint=os.getenv("DOCUMENTINTELLIGENCE_ENDPOINT"),
    credential=AzureKeyCredential(os.getenv("DOCUMENTINTELLIGENCE_API_KEY"))
)

deployment_name = 'gpt-4o'


def process_financial_report(pdf_path):
    """Extracts financial data from a PDF using Azure Document Intelligence."""

    with open(pdf_path, "rb") as f:
        poller = document_intelligence_client.begin_analyze_document(
            model_id="prebuilt-layout",
            analyze_request=f,
        )
        result = poller.result()

    extracted_data = {}
    for page in result.pages:
        for table in page.tables:
            for cell in table.cells:
                extracted_data[cell.content] = cell.bounding_region[0].polygon if cell.bounding_region else None
    return extracted_data

def calculate_ratios(data):
    """Calculates financial ratios from extracted financial data."""
    total_assets = float(data.get("Total Assets", 1))
    total_liabilities = float(data.get("Total Liabilities", 1))
    shareholder_equity = float(data.get("Shareholder Equity", 1))
    net_income = float(data.get("Net Income", 1))
    total_revenue = float(data.get("Total Revenue", 1))
    earnings_per_share = float(data.get("EPS", 1))

    ratios = {
        "Current Ratio": total_assets / total_liabilities,
        "Debt-to-Equity Ratio": total_liabilities / shareholder_equity,
        "Return on Equity (ROE)": (net_income / shareholder_equity) * 100,
        "Return on Assets (ROA)": (net_income / total_assets) * 100,
        "Gross Profit Margin": (net_income / total_revenue) * 100,
        "Net Profit Margin": (net_income / total_revenue) * 100,
        "Earnings Per Share (EPS)": earnings_per_share
    }
    return ratios

def generate_summary(data, ratios):
    """Creates a one-page financial summary."""
    prompt = f"""
    Summarize the following financial performance data:
    {data}
    Include risk factors and overall performance highlights.
    """
    response = azure_openai_client.Completion.create(
        engine=deployment_name,
        prompt=prompt,
        max_tokens=300,
        temperature=0.7
    )
    return response["choices"][0]["text"]

def create_excel(ratios, output_path="financial_ratios.xlsx"):
    """Saves the calculated ratios into an Excel file."""
    df = pd.DataFrame([ratios])
    df.to_excel(output_path, index=False)

def create_word_document(summary, output_path="financial_summary.docx"):
    """Generates a Word document with the performance summary."""
    doc = Document()
    doc.add_heading("Financial Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(output_path)

if __name__ == "__main__":
    pdf_path = "Report.pdf"

    extracted_data = process_financial_report(pdf_path)

    ratios = calculate_ratios(extracted_data)

    summary = generate_summary(extracted_data, ratios)

    create_excel(ratios)
    create_word_document(summary)

    print("Financial analysis completed. Files generated: 'financial_ratios.xlsx', 'financial_summary.docx'")
