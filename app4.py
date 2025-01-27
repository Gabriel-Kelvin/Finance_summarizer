import os
import pandas as pd
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from docx import Document

# Azure Form Recognizer credentials
FORM_RECOGNIZER_ENDPOINT = "https://neohack.cognitiveservices.azure.com/"
FORM_RECOGNIZER_API_KEY = "D4Vd1Xljol6cS0q8PoZiQFXz8uH1yYeXCIyhTHU5vdBxBhbeX9JLJQQJ99ALACYeBjFXJ3w3AAALACOGIiVt"

# Initialize Azure Document Analysis Client
document_analysis_client = DocumentAnalysisClient(
    endpoint=FORM_RECOGNIZER_ENDPOINT,
    credential=AzureKeyCredential(FORM_RECOGNIZER_API_KEY)
)


def process_financial_report(pdf_path):
    """
    Extracts financial data from a PDF using Azure Document Analysis Client.
    """
    try:
        with open(pdf_path, "rb") as file:
            # Analyze the document using the prebuilt-layout model
            poller = document_analysis_client.begin_analyze_document(
                model_id="prebuilt-layout",
                document=file
            )
            result = poller.result()

        # Extract tables from the document
        extracted_data = {}
        for table in result.tables:
            for cell in table.cells:
                extracted_data[cell.content.strip()] = (
                    cell.bounding_regions[0].polygon if cell.bounding_regions else None
                )
        return extracted_data

    except Exception as e:
        raise RuntimeError(f"Error processing financial report: {e}")


def calculate_ratios(data):
    """
    Calculates financial ratios from extracted financial data.
    """
    try:
        total_assets = float(data.get("Total Assets", 1))
        total_liabilities = float(data.get("Total Liabilities", 1))
        shareholder_equity = float(data.get("Shareholder Equity", 1))
        net_income = float(data.get("Net Income", 1))
        total_revenue = float(data.get("Total Revenue", 1))

        ratios = {
            "Current Ratio": total_assets / total_liabilities,
            "Debt-to-Equity Ratio": total_liabilities / shareholder_equity,
            "Return on Equity (ROE)": (net_income / shareholder_equity) * 100,
            "Return on Assets (ROA)": (net_income / total_assets) * 100,
            "Gross Profit Margin": (net_income / total_revenue) * 100,
            "Net Profit Margin": (net_income / total_revenue) * 100,
        }
        return ratios

    except Exception as e:
        raise ValueError(f"Error calculating financial ratios: {e}")


def create_excel(ratios, output_path="financial_ratios.xlsx"):
    """
    Saves the calculated ratios into an Excel file.
    """
    try:
        df = pd.DataFrame([ratios])
        df.to_excel(output_path, index=False)
        print(f"Excel file saved at {output_path}")
    except Exception as e:
        raise RuntimeError(f"Error creating Excel file: {e}")


def create_word_document(data, ratios, output_path="financial_summary.docx"):
    """
    Generates a Word document with financial data and calculated ratios.
    """
    try:
        doc = Document()
        doc.add_heading("Financial Summary", level=1)

        doc.add_heading("Extracted Financial Data", level=2)
        for key, value in data.items():
            doc.add_paragraph(f"{key}: {value}")

        doc.add_heading("Calculated Financial Ratios", level=2)
        for key, value in ratios.items():
            doc.add_paragraph(f"{key}: {value:.2f}")

        doc.save(output_path)
        print(f"Word document saved at {output_path}")
    except Exception as e:
        raise RuntimeError(f"Error creating Word document: {e}")


if __name__ == "__main__":
    pdf_path = "Report.pdf"

    try:
        extracted_data = process_financial_report(pdf_path)
        print("Financial data extracted successfully!")

        ratios = calculate_ratios(extracted_data)
        print("Financial ratios calculated successfully!")

        create_excel(ratios)

        create_word_document(extracted_data, ratios)

        print("Financial analysis completed successfully!")

    except Exception as e:
        print(f"Error: {e}")
